from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import os
from datetime import datetime

# -------------------------- Word生成函数 --------------------------
def generate_doubao_chat_word(messages, title="豆包大模型对话记录"):
    """从豆包对话历史生成Word文档，返回文件路径+文件名"""
    doc = Document()
    # 设置标题：居中、黑体、二号字、加粗
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.name = "黑体"
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    title_run.font.size = Pt(22)
    title_run.bold = True
    title_para.alignment = 1  # 1=居中

    # 添加生成时间：宋体、小四、右对齐
    time_para = doc.add_paragraph()
    time_str = f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}"
    time_run = time_para.add_run(time_str)
    time_run.font.name = "宋体"
    time_run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    time_run.font.size = Pt(12)
    time_para.alignment = 2  # 2=右对齐

    # 添加强制分页符
    doc.add_page_break()

    # 遍历对话历史，写入Word（区分用户/豆包，自定义样式）
    for msg in messages:
        role = msg["role"]
        content = msg["content"].strip()
        if not content:
            continue

        para = doc.add_paragraph()
        # 角色样式：用户蓝色、豆包深灰，均加粗
        if role == "user":
            role_run = para.add_run("用户：")
            role_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)  # 蓝色
        else:
            role_run = para.add_run("豆包：")
        role_run.font.name = "宋体"
        role_run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        role_run.font.size = Pt(12)
        role_run.bold = True

        # 对话内容：宋体、小四、首行缩进2字符、1.5倍行间距
        content_run = para.add_run(content)
        content_run.font.name = "宋体"
        content_run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        content_run.font.size = Pt(12)
        para.paragraph_format.first_line_indent = Inches(0.29)  # 2字符=0.29英寸
        para.paragraph_format.line_spacing = 1.5  # 1.5倍行间距

    # 生成唯一文件名，避免重复
    file_name = f"豆包对话_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    file_path = os.path.join(os.path.dirname(__file__), "static/generate_files/word", file_name)
    # 保存文档
    doc.save(file_path)
    return file_path, file_name

# -------------------------- HTML生成函数 --------------------------
def generate_doubao_chat_html(messages, title="豆包大模型对话记录", custom_css=None):
    """从豆包对话历史生成HTML文件，返回文件路径+文件名+纯HTML代码"""
    # 默认CSS样式（响应式，适配电脑/手机，美观简洁）
    default_css = """
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; font-family: "Microsoft YaHei", "PingFang SC", sans-serif; }
        body { max-width: 1200px; margin: 20px auto; padding: 0 20px; background-color: #f5f7fa; }
        .page-title { text-align: center; font-size: 24px; font-weight: 700; color: #2d3748; margin-bottom: 20px; }
        .generate-time { text-align: right; font-size: 14px; color: #718096; margin-bottom: 30px; }
        .chat-container { display: flex; flex-direction: column; gap: 16px; margin-top: 20px; }
        .chat-item { padding: 12px 16px; border-radius: 8px; max-width: 90%; word-wrap: break-word; }
        .user-chat { background-color: #e6f7ff; align-self: flex-end; border: 1px solid #91d5ff; }
        .doubao-chat { background-color: #ffffff; align-self: flex-start; border: 1px solid #e2e8f0; }
        .chat-role { font-weight: 700; margin-right: 8px; }
        .user-role { color: #1890ff; }
        .doubao-role { color: #2d3748; }
        .chat-content { font-size: 16px; line-height: 1.6; color: #2d3748; }
    </style>
    """
    # 使用自定义CSS（若有），否则用默认样式
    css_style = custom_css if custom_css and isinstance(custom_css, str) else default_css

    # 拼接HTML头部（含meta、标题、样式）
    html_head = f"""
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{title}</title>
        {css_style}
    </head>
    <body>
        <h1 class="page-title">{title}</h1>
        <div class="generate-time">生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}</div>
        <div class="chat-container">
    """

    # 拼接对话记录（区分用户/豆包，绑定不同样式类）
    html_chat_content = ""
    for msg in messages:
        role = msg["role"]
        content = msg["content"].strip().replace("\n", "<br>")  # 换行符转HTML<br>标签
        if not content:
            continue
        if role == "user":
            html_chat_content += f"""
            <div class="chat-item user-chat">
                <span class="chat-role user-role">用户：</span>
                <span class="chat-content">{content}</span>
            </div>
            """
        else:
            html_chat_content += f"""
            <div class="chat-item doubao-chat">
                <span class="chat-role doubao-role">豆包：</span>
                <span class="chat-content">{content}</span>
            </div>
            """

    # 拼接HTML尾部
    html_foot = """
        </div>
    </body>
    </html>
    """

    # 拼接完整HTML代码
    full_html = html_head + html_chat_content + html_foot

    # 生成唯一文件名，保存HTML文件
    file_name = f"豆包对话_{datetime.now().strftime('%Y%m%d%H%M%S')}.html"
    file_path = os.path.join(os.path.dirname(__file__), "static/generate_files/html", file_name)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(full_html)

    return file_path, file_name, full_html

# -------------------------- 图像生成函数 --------------------------
def generate_doubao_chat_image(messages, title="豆包大模型对话记录", img_size=(1200, 800), img_format="PNG"):
    """从豆包对话历史生成图像，返回文件路径+文件名，支持PNG/JPG，自定义尺寸"""
    # 初始化画布：白色背景、RGB模式
    img = Image.new("RGB", img_size, color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    # 配置字体（兼容Windows/Linux/Mac，优先系统中文字体，兜底默认字体）
    try:
        # Windows系统默认中文字体
        font_title = ImageFont.truetype("simhei.ttf", 36)  # 标题：黑体36号
        font_role = ImageFont.truetype("simsun.ttc", 24)   # 角色：宋体24号（加粗）
        font_content = ImageFont.truetype("simsun.ttc", 20)# 内容：宋体20号
    except:
        try:
            # Linux/Mac开源中文字体（Noto Sans CJK）
            font_title = ImageFont.truetype("NotoSansCJK-Regular.ttc", 36)
            font_role = ImageFont.truetype("NotoSansCJK-Regular.ttc", 24)
            font_content = ImageFont.truetype("NotoSansCJK-Regular.ttc", 20)
        except:
            # 兜底：使用Pillow默认字体，确保文字正常显示
            font_title = ImageFont.load_default(size=36)
            font_role = ImageFont.load_default(size=24)
            font_content = ImageFont.load_default(size=20)

    # 绘制配置：坐标、颜色、行高
    x_start, y_start = 50, 30  # 画布起始绘制坐标
    line_height = 35  # 文字行高
    # 角色颜色：用户蓝色（0,0,255），豆包深灰（51,51,51）
    role_colors = {"user": (0, 0, 255), "assistant": (51, 51, 51)}
    title_color = (26, 26, 26)  # 标题颜色：深黑色
    time_color = (113, 128, 150)  # 时间颜色：灰色

    # 绘制标题（居中显示）
    title_width, title_height = draw.textsize(title, font=font_title)
    title_x = (img_size[0] - title_width) // 2  # 居中X坐标
    draw.text((title_x, y_start), title, fill=title_color, font=font_title)
    y_current = y_start + title_height + 40  # 更新当前绘制Y坐标（标题下方留40px间距）

    # 绘制生成时间
    time_str = f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}"
    draw.text((x_start, y_current), time_str, fill=time_color, font=font_content)
    y_current += line_height + 20  # 时间下方留20px间距

    # 绘制对话记录（自动换行，避免超出画布边界）
    max_content_width = img_size[0] - 100  # 内容最大宽度（左右各留50px边距）
    for msg in messages:
        role = msg["role"]
        content = msg["content"].strip()
        # 跳过空内容，或画布剩余高度不足时停止绘制
        if not content or y_current + 100 > img_size[1]:
            break

        # 绘制角色标签（如：用户：/ 豆包：）
        role_text = "用户：" if role == "user" else "豆包："
        draw.text((x_start, y_current), role_text, fill=role_colors[role], font=font_role)
        # 计算角色标签宽度，确定内容起始X坐标
        role_width, _ = draw.textsize(role_text, font=font_role)
        content_x = x_start + role_width + 10  # 角色右侧留10px间距

        # 内容自动换行处理：逐字符判断，超出最大宽度则换行
        current_line = ""
        content_lines = []
        for char in content:
            line_width, _ = draw.textsize(current_line + char, font=font_content)
            if line_width > max_content_width - content_x:
                content_lines.append(current_line)
                current_line = char
            else:
                current_line += char
        if current_line:
            content_lines.append(current_line)

        # 逐行绘制对话内容
        for line in content_lines:
            if y_current + line_height > img_size[1]:
                break
            draw.text((content_x, y_current), line, fill=role_colors[role], font=font_content)
            y_current += line_height
        y_current += 10  # 对话项之间留10px间距

    # 处理图像格式，生成唯一文件名
    img_format = img_format.upper()
    file_suffix = "png" if img_format == "PNG" else "jpg"
    file_name = f"豆包对话_{datetime.now().strftime('%Y%m%d%H%M%S')}.{file_suffix}"
    file_path = os.path.join(os.path.dirname(__file__), "static/generate_files/image", file_name)

    # 保存图像：JPG格式需去除透明通道，设置质量
    if img_format == "JPG":
        img.save(file_path, format="JPEG", quality=95)
    else:
        img.save(file_path, format="PNG")

    return file_path, file_name